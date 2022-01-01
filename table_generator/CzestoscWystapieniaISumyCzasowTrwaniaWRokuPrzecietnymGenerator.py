from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from Constants import PROVIDED_INVALID_WATER_PARAMETER_NAME_ERROR_MESSAGE
from table_generator.TableGenerator import addHeadersToTable, SIGMA_SIGN, addStyledContentToCellAndMerge


def prepareHeadersForCzestoscWystapieniaISumyCzasowTrwania(avg_years, parameter_name):
    if parameter_name == 'h_water':
        przedzial_stanow_header_text = 'Przedział stanów [cm]'
        czestosc_wystapienia_stanow_header_text = 'Częstość wystąpienia stanów [cm]'
    elif parameter_name == 'Q':
        przedzial_stanow_header_text = 'Przedział przepływów [m\u00B3/s]'
        czestosc_wystapienia_stanow_header_text = 'Częstość wystąpienia przepływów [m\u00B3/s]'
    else:
        raise ValueError(PROVIDED_INVALID_WATER_PARAMETER_NAME_ERROR_MESSAGE)

    przedzial_stanow_header = (przedzial_stanow_header_text, 0, 0, 3, 4)
    czestosc_wystapienia_stanow_header = (czestosc_wystapienia_stanow_header_text, 0, 4, 2, 8)
    frequency_header_header = ('Suma częstości [dni]', 0, 12, 3, 3)
    average_year_header = ('Rok przeciętny [dni]', 0, 15, 3, 3)
    sumy_czasow_trwania_header = ('Sumy czasów trwania wraz z', 0, 18, 2, 4)
    lower_header_header = ('niższymi', 2, 18, 1, 2)
    higher_header_text = ('wyższymi', 2, 20, 1, 2)

    years_range_headers = []
    column_index_for_year_header_cell = 4
    for year_number in avg_years:
        years_range_headers.append((str(year_number), 2, column_index_for_year_header_cell, 1, 2))
        column_index_for_year_header_cell += 2

    headers_definitions = [przedzial_stanow_header,
                           czestosc_wystapienia_stanow_header,
                           frequency_header_header, average_year_header,
                           sumy_czasow_trwania_header]
    headers_definitions.extend(years_range_headers)
    headers_definitions.extend([lower_header_header, higher_header_text])
    return headers_definitions


def appendCzestoscWystapieniaISumyCzasowTrwaniaWRokuPrzecietnymTableToDocument(document, states_in_average_years_dataFrame, years_range, parameter_name):
    number_of_columns = 22
    headers_for_czestosc_wystapienia_i_sumy_czasow_trwania_table = prepareHeadersForCzestoscWystapieniaISumyCzasowTrwania(years_range, parameter_name)
    headers_for_czestosc_wystapienia_i_sumy_czasow_trwania_table.sort(key=lambda header: header[1], reverse=True)

    number_of_rows_for_headers_cells = headers_for_czestosc_wystapienia_i_sumy_czasow_trwania_table[0][1] + 1
    average_year_table = document.add_table(rows=number_of_rows_for_headers_cells, cols=number_of_columns)
    average_year_table.style = 'Table Grid'

    addHeadersToTable(average_year_table, headers_for_czestosc_wystapienia_i_sumy_czasow_trwania_table)
    __addTableContent(average_year_table, states_in_average_years_dataFrame, years_range)


def __addTableContent(table, table_content, years_range):
    for state_range, table_data_row in table_content.iterrows():
        row_cells = table.add_row().cells
        state_cell = row_cells[0]
        paragraph = state_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.style.font.size = Pt(8)
        run = paragraph.add_run(str(state_range))
        run.bold = True
        state_cell.merge(row_cells[1])
        state_cell.merge(row_cells[2])
        state_cell.merge(row_cells[3])

        cell_column_index = 4
        for current_year in years_range:
            addStyledContentToCellAndMerge(row_cells, cell_column_index, table_data_row[str(current_year)], 1)
            row_cells[cell_column_index].merge(row_cells[cell_column_index + 1])
            cell_column_index += 2

        frequency_first_cell_index = cell_column_index
        addStyledContentToCellAndMerge(row_cells, frequency_first_cell_index, table_data_row['frequency'], 3)

        average_year_first_cell_index = frequency_first_cell_index + 3
        addStyledContentToCellAndMerge(row_cells, average_year_first_cell_index, table_data_row['average year'], 3)

        lower_first_cell_index = average_year_first_cell_index + 3
        addStyledContentToCellAndMerge(row_cells, lower_first_cell_index, table_data_row['lower'], 2)

        higher_first_cell_index = lower_first_cell_index + 2
        addStyledContentToCellAndMerge(row_cells, higher_first_cell_index, table_data_row['higher'], 2)

        if state_range == 'sum':
            __formatLastRow(row_cells, frequency_first_cell_index, lower_first_cell_index, higher_first_cell_index)


def __formatLastRow(row_cells, frequency_first_cell_index, lower_first_cell_index, higher_first_cell_index):
    run_with_text = row_cells[0].paragraphs[0].runs[0]
    run_with_text.clear()
    run_with_sigma = row_cells[0].paragraphs[0].add_run(SIGMA_SIGN)
    run_with_sigma.bold = True

    run_with_frequency_value = row_cells[frequency_first_cell_index].paragraphs[0].runs[0]
    run_with_frequency_value.clear()

    run_with_lower_value = row_cells[lower_first_cell_index].paragraphs[0].runs[0]
    run_with_lower_value.clear()

    run_with_higher_value = row_cells[higher_first_cell_index].paragraphs[0].runs[0]
    run_with_higher_value.clear()
