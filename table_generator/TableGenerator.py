from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL

SIGMA_SIGN = '\u03A3'


def __fillAndMergeCells(table, cell_content, start_row_position, start_column_position, row_span, column_span):
    parent_cell = table.cell(start_row_position, start_column_position)
    paragraph = parent_cell.paragraphs[0]
    parent_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.style.font.size = Pt(8)
    run = paragraph.add_run(cell_content)
    run.bold = True
    for row_idx in range(start_row_position, start_row_position + row_span):
        parent_cell = table.cell(row_idx, start_column_position)
        for column_idx in range(start_column_position, start_column_position + column_span):
            parent_cell.merge(table.cell(row_idx, column_idx))
        if row_idx != start_row_position:
            parent_cell.merge(table.cell(row_idx - 1, start_column_position))


def addHeadersToTable(table, headers_definitions):
    for cell_content, start_row_position, start_column_position, row_span, column_span in headers_definitions:
        __fillAndMergeCells(table, cell_content, start_row_position, start_column_position, row_span, column_span)


def addStyledContentToCellAndMerge(row_cells, parent_cell_index, cell_content, merged_cell_length):
    paragraph = row_cells[parent_cell_index].paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.style.font.size = Pt(8)
    paragraph.add_run(str(cell_content))

    first_not_appended_cell_index = parent_cell_index + merged_cell_length
    for cell_index_to_append in range(parent_cell_index + 1, first_not_appended_cell_index):
        row_cells[parent_cell_index].merge(row_cells[cell_index_to_append])
    return first_not_appended_cell_index
