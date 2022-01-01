from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from HydrologyReportCreator import ORDERED_HYDRO_MONTHS, PROVIDED_INVALID_WATER_PARAMETER_NAME_ERROR_MESSAGE
from table_generator.TableGenerator import addHeadersToTable, SIGMA_SIGN, addStyledContentToCellAndMerge


def prepareHeadersForCzestoscWystapieniaISumyCzasowTrwania(parameter_name):
    if parameter_name == 'h_water':
        przedzial_stanow_header = ('Przedział stanów [cm]', 0, 0, 3, 3)
    elif parameter_name == 'Q':
        przedzial_stanow_header = ('Przedział przepływów [m\u00B3/s]', 0, 0, 3, 3)
    else:
        raise ValueError(PROVIDED_INVALID_WATER_PARAMETER_NAME_ERROR_MESSAGE)

    czestosc_wystapienia_stanow_header = ('Częstość wystąpienia stanów [dni]', 0, 3, 2, 12)
    frequency_header_header = ('Częstość [dni]', 0, 15, 3, 3)
    sumy_czasow_trwania_header = ('Sumy czasów trwania wraz z', 0, 18, 2, 4)
    lower_header_header = ('niższymi', 2, 18, 1, 2)
    higher_header_text = ('wyższymi', 2, 20, 1, 2)

    months_headers = []
    column_index_for_month_header_cell = 3
    for month_number in ORDERED_HYDRO_MONTHS:
        months_headers.append((str(month_number), 2, column_index_for_month_header_cell, 1, 1))
        column_index_for_month_header_cell += 1

    headers_definitions = [przedzial_stanow_header,
                           czestosc_wystapienia_stanow_header,
                           frequency_header_header,
                           sumy_czasow_trwania_header]
    headers_definitions.extend(months_headers)
    headers_definitions.extend([lower_header_header, higher_header_text])
    return headers_definitions


def __addTableContent(table, table_content):
    for state_range, table_data_row in table_content.iterrows():
        row_cells = table.add_row().cells
        state_cell = row_cells[0]
        paragraph = state_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.style.font.size = Pt(8)
        run = paragraph.add_run(str(state_range))
        run.bold = True
        state_cell.merge(row_cells[1])
        state_cell.merge(row_cells[2])

        cell_column_index = 3
        for current_month in ORDERED_HYDRO_MONTHS:
            addStyledContentToCellAndMerge(row_cells, cell_column_index, table_data_row[str(current_month)], 1)
            cell_column_index += 1

        addStyledContentToCellAndMerge(row_cells, 15, table_data_row['frequency'], 3)
        addStyledContentToCellAndMerge(row_cells, 18, table_data_row['lower'], 2)
        addStyledContentToCellAndMerge(row_cells, 20, table_data_row['higher'], 2)

        if state_range == 'sum':
            __formatLastRow(row_cells)


def __formatLastRow(row_cells):
    run_with_text = row_cells[0].paragraphs[0].runs[0]
    run_with_text.clear()
    run_with_sigma = row_cells[0].paragraphs[0].add_run(SIGMA_SIGN)
    run_with_sigma.bold = True

    run_with_lower_value = row_cells[18].paragraphs[0].runs[0]
    run_with_lower_value.clear()

    run_with_higher_value = row_cells[20].paragraphs[0].runs[0]
    run_with_higher_value.clear()


def appendCzestoscWystapieniaISumyCzasowTrwaniaTableToDocument(document, table_content, parameter_name):
    number_of_columns = 22
    headers_for_czestosc_wystapienia_i_sumy_czasow_trwania_table = prepareHeadersForCzestoscWystapieniaISumyCzasowTrwania(parameter_name)
    headers_for_czestosc_wystapienia_i_sumy_czasow_trwania_table.sort(key=lambda header: header[1], reverse=True)

    number_of_rows_for_headers_cells = headers_for_czestosc_wystapienia_i_sumy_czasow_trwania_table[0][1] + 1
    czestoscWystapieniaISumyCzasowTrwaniaTable = document.add_table(rows=number_of_rows_for_headers_cells, cols=number_of_columns)
    czestoscWystapieniaISumyCzasowTrwaniaTable.style = 'Table Grid'
    czestoscWystapieniaISumyCzasowTrwaniaTable.allow_autofit = True

    addHeadersToTable(czestoscWystapieniaISumyCzasowTrwaniaTable, headers_for_czestosc_wystapienia_i_sumy_czasow_trwania_table)
    __addTableContent(czestoscWystapieniaISumyCzasowTrwaniaTable, table_content)
