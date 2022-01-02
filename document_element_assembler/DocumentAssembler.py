import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.shared import Pt

from Constants import TEMP_FOLDER_DIRECTORY, CHART_IMAGES_DIRECTORY
from chart_generator import MainStatesFluctuationCurveForYears


def createDocumentWithFirstPage(city_name, river_name):
    document = Document()
    __appendFirstPageToDocument(document, city_name, river_name)
    return document


def __appendFirstPageToDocument(document, city_name, river_name):
    run = document.add_paragraph().add_run()
    font = run.font

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph1 = document.add_paragraph('Katedra Hydrotechniki\nWydział Inżynierii Lądowej i Środowiska\nPolitechniki Gdańskiej')

    font.name = 'Arial'
    font.size = Pt(12)
    paragraph1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_picture('./resources/pg_logo.png', width=Cm(14))
    title = document.add_heading('OPERAT HYDROLOGICZNY', 0)
    title.add_run().bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    riverAndCityParagraphContent = 'RZEKA:%s \nPRZEKRÓJ:%s' % (river_name, city_name)
    r1 = document.add_paragraph(riverAndCityParagraphContent)
    r1.add_run().bold = True
    r1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def addChartToDocument(document, page_heading_content, chart_filename):
    document.add_page_break()
    page_heading = document.add_heading(page_heading_content, 2)
    page_heading.add_run().bold = True
    page_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    path_to_chart_file = os.path.join(TEMP_FOLDER_DIRECTORY, CHART_IMAGES_DIRECTORY, chart_filename)
    document.add_picture(path_to_chart_file, width=Cm(19))


def addMainStatesFluctuationCurveForYears(document, index_of_element, main_states_first_degree,
                                          main_states_second_degree,
                                          river_name, city_name,
                                          year_from, year_to):
    krzywa_wahan_stanow_glownych_1_stopnia_chart_filename, krzywa_wahan_stanow_glownych_1_stopnia_chart_title = MainStatesFluctuationCurveForYears.printMainStatesFluctuationCurveForYears(
        main_states_first_degree,
        main_states_second_degree,
        river_name, city_name,
        year_from, year_to)

    page_heading_content = '%d. %s' % (index_of_element, krzywa_wahan_stanow_glownych_1_stopnia_chart_title)
    addChartToDocument(document, page_heading_content, krzywa_wahan_stanow_glownych_1_stopnia_chart_filename)
