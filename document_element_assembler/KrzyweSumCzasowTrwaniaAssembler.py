from docx import Document

import Constants
from Constants import WATER_PARAMETERS
from chart_generator import HistogramCzestosciWystapieniaWRokuPrzecietnym, KrzyweSumCzasowTrwania
from dataset_provider import DatasetProvider, DataFrameForCzestoscWystapieniaISumyCzasowTrwaniaWRokuPrzecietnymProvider
from dataset_provider.DataFrameForCzestoscWystapieniaISumyCzasowTrwaniaProvider import createDataFrameForCzestoscWystapieniaISumyCzasowTrwaniaTable
from document_element_assembler import DocumentAssembler
from table_generator import CzestoscWystapieniaISumyCzasowTrwaniaTableGenerator, CzestoscWystapieniaISumyCzasowTrwaniaWRokuPrzecietnymGenerator


def addHistogramsAndCzestoscWystapieniaAndSumyCzasowTrwaniaElements(document: Document, index_of_element, main_states_first_degree,
                                                                    river_name, city_name, first_year_of_multiannual_period, dataset_for_years,
                                                                    include_first_row_for_higher_in_frequency_table, common_range_of_states_for_multiannual):
    inner_index_of_element = 1
    year_to_multiannual = first_year_of_multiannual_period + Constants.MULTIANNUAL_PERIOD - 1  # -1 as we also include first year of multiannual period
    multiannual_years_range = range(first_year_of_multiannual_period, year_to_multiannual + 1)
    wet_year, dry_year = __computeDryAndWetYears(main_states_first_degree, multiannual_years_range)

    for parameter_name in WATER_PARAMETERS:
        min_value, max_value = DatasetProvider.findMinAndMaxValueOfWaterParameterInYearsRange(dataset_for_years, multiannual_years_range, parameter_name)
        states = DatasetProvider.createStatesForWaterParameter(min_value, max_value, parameter_name)

        dailyFlowsAndStatesInYears = createDataFrameForCzestoscWystapieniaISumyCzasowTrwaniaTable(dataset_for_years, multiannual_years_range, parameter_name,
                                                                                                  include_first_row_for_higher_in_frequency_table,
                                                                                                  common_range_of_states_for_multiannual, False)
        states_in_average_years = DataFrameForCzestoscWystapieniaISumyCzasowTrwaniaWRokuPrzecietnymProvider.createDataFrameForCzestoscWystapieniaISumyCzasowTrwaniaWRokuPrzecietnymTable(
            dataset_for_years, multiannual_years_range, parameter_name, include_first_row_for_higher_in_frequency_table, False)

        for current_year in [dry_year, wet_year]:
            year_type = 'dry' if current_year == dry_year else 'wet'
            inner_index_of_element = breakPageAndAddHeading(document, year_type, parameter_name, current_year, current_year, river_name, city_name,
                                                            index_of_element, inner_index_of_element, 'table')
            current_year_df_for_table = dailyFlowsAndStatesInYears[current_year]
            CzestoscWystapieniaISumyCzasowTrwaniaTableGenerator.appendCzestoscWystapieniaISumyCzasowTrwaniaTableToDocument(document, current_year_df_for_table,
                                                                                                                           parameter_name)
        # add table for average year
        inner_index_of_element = breakPageAndAddHeading(document, 'average', parameter_name, first_year_of_multiannual_period, year_to_multiannual, river_name,
                                                        city_name, index_of_element, inner_index_of_element, 'table')
        CzestoscWystapieniaISumyCzasowTrwaniaWRokuPrzecietnymGenerator.appendCzestoscWystapieniaISumyCzasowTrwaniaWRokuPrzecietnymTableToDocument(document,
                                                                                                                                                  states_in_average_years,
                                                                                                                                                  multiannual_years_range,
                                                                                                                                                  parameter_name)

        # add histogram
        histogram_file_name, histogram_chart_title = HistogramCzestosciWystapieniaWRokuPrzecietnym.printHistogramCzestosciWystapieniaWRokuPrzecietnym(
            states_in_average_years, states,
            parameter_name,
            multiannual_years_range[0],
            multiannual_years_range[-1], city_name,
            river_name)
        inner_index_of_element = breakPageAndAddHeading(document, 'average', parameter_name, first_year_of_multiannual_period, year_to_multiannual, river_name,
                                                        city_name, index_of_element, inner_index_of_element, 'histogram')
        DocumentAssembler.addChartToDocument(document, histogram_file_name)

        # add Krzywe sum czasów trwania
        wet_year_df = dailyFlowsAndStatesInYears[wet_year]
        dry_year_df = dailyFlowsAndStatesInYears[dry_year]

        inner_index_of_element = breakPageAndAddHeading(document, 'average', parameter_name, first_year_of_multiannual_period, year_to_multiannual, river_name,
                                                        city_name, index_of_element, inner_index_of_element, 'krzywe_sum_czasow_trwania')
        krzywe_sum_czasow_trwania_file_name, _ = KrzyweSumCzasowTrwania.createChart(dry_year_df, wet_year_df, states_in_average_years, parameter_name, dry_year,
                                                                                    wet_year, city_name, river_name)
        DocumentAssembler.addChartToDocument(document, krzywe_sum_czasow_trwania_file_name)


def breakPageAndAddHeading(document, year_type, parameter_name, year_from, year_to, river_name, city_name, index_of_element, inner_index_of_element,
                           element_type):
    document.add_page_break()
    textual_element_index = '%d.%d' % (index_of_element, inner_index_of_element)
    heading_for_element = __provideHeadingForElement(year_type, parameter_name, year_from, year_to, element_type, river_name, city_name)
    DocumentAssembler.addHeadingToDocumentElement(document, textual_element_index, heading_for_element)
    inner_index_of_element += 1
    return inner_index_of_element


def __computeDryAndWetYears(main_states_first_degree, multiannual_years_range):
    main_states_for_multiannual = main_states_first_degree[main_states_first_degree.index.isin(multiannual_years_range)]
    wet_year = main_states_for_multiannual['SW'].idxmax()
    dry_year = main_states_for_multiannual['SW'].idxmin()
    return wet_year, dry_year


def __provideHeadingForElement(year_type, parameter_name, year_from, year_to, element_type, river_name, city_name):
    if parameter_name == 'h_water':
        parameter_name_text = 'stanów'
    else:
        parameter_name_text = 'przepływów'

    if year_type == 'wet':
        year_type_text = 'mokrym'
    elif year_type == 'dry':
        year_type_text = 'suchym'
    else:
        year_type_text = 'przeciętnym'

    city_river_text = f'\nrzeka:{river_name}, przekrój:{city_name}'
    if element_type == 'krzywe_sum_czasow_trwania':
        return f'Krzywe sum czasów trwania {parameter_name_text} wraz z wyższymi' + city_river_text
    if element_type == 'histogram':
        return f'Histogram częstości wystąpienia {parameter_name_text} w roku przeciętnym' + city_river_text
    if element_type == 'table' and year_type == 'average':
        return f'Częstość wystąpienia i sumy czasów trwania {parameter_name_text} w roku przeciętnym dla okresu {year_from}-{year_to}' + city_river_text
    if element_type == 'table':
        return f'Częstość wystąpienia i sumy czasów trwania {parameter_name_text} w roku {year_type_text} {year_from}' + city_river_text
    return ''
